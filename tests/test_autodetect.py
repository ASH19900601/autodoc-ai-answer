"""Tests for auto mode detection (simple vs worksheet)."""
from __future__ import annotations

import json
from pathlib import Path

import docx
import pytest

from autodoc.core.answer_router import answer as route_answer
from autodoc.core.answer_router import choose_mode


class FakeOpsLLM:
    """Returns worksheet ops; records that it was called."""

    def __init__(self):
        self.called = False

    def complete(self, system, user):
        self.called = True
        return json.dumps([{"op": "set_paragraph", "index": 2, "text": "ans"}])


class FakeMapLLM:
    """Returns a simple-mode answer map."""

    def complete(self, system, user):
        return json.dumps({"q1": "北京", "q2": "H2O", "q3": "月球"})


def test_auto_picks_simple_for_numbered(sample_docx: Path):
    # sample_docx: heading + numbered fill-blank questions, no tables.
    assert choose_mode(str(sample_docx)) == "simple"


def test_auto_picks_worksheet_when_tables(tmp_path: Path):
    d = docx.Document()
    d.add_paragraph("(a) Explain something.")
    d.add_paragraph("")
    t = d.add_table(rows=2, cols=3)
    t.rows[0].cells[1].text = "True"
    t.rows[0].cells[2].text = "False"
    t.rows[1].cells[0].text = "A statement."
    p = tmp_path / "ws.docx"
    d.save(str(p))
    assert choose_mode(str(p)) == "worksheet"


def test_auto_picks_worksheet_for_lettered_parts(tmp_path: Path):
    d = docx.Document()
    d.add_paragraph("(a) State one purpose of a router. [1]")
    d.add_paragraph("")
    d.add_paragraph("(b) Explain IP assignment. [2]")
    d.add_paragraph("")
    p = tmp_path / "lettered.docx"
    d.save(str(p))
    assert choose_mode(str(p)) == "worksheet"


def test_non_docx_is_simple(sample_txt: Path):
    assert choose_mode(str(sample_txt)) == "simple"


def test_route_auto_simple_sets_mode(sample_docx: Path):
    result = route_answer(str(sample_docx), FakeMapLLM(), mode="auto")
    assert result.mode == "simple"
    assert result.answers_written == 3


def test_route_auto_worksheet_sets_mode(tmp_path: Path):
    d = docx.Document()
    d.add_paragraph("(a) Explain something. [1]")
    d.add_paragraph("")
    t = d.add_table(rows=2, cols=3)
    t.rows[0].cells[1].text = "True"
    t.rows[0].cells[2].text = "False"
    t.rows[1].cells[0].text = "A statement."
    p = tmp_path / "ws.docx"
    d.save(str(p))

    llm = FakeOpsLLM()
    result = route_answer(str(p), llm, mode="auto")
    assert result.mode == "worksheet"
    assert llm.called is True


def test_route_explicit_mode_overrides(sample_docx: Path):
    result = route_answer(str(sample_docx), FakeMapLLM(), mode="simple")
    assert result.mode == "simple"
