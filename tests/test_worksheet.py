"""Worksheet mode tests: structure extraction + in-place op application.

Deterministic (no LLM): mirrors the real GCSE-style worksheet shape -
lettered sub-parts, blank answer paragraphs, and a True/False tick table.
"""
from __future__ import annotations

import zipfile
from pathlib import Path

import docx
import pytest
from docx.shared import Pt

from autodoc.core.backends import docx_backend
from autodoc.core.worksheet import extract_json_array


@pytest.fixture
def worksheet_docx(tmp_path: Path) -> Path:
    d = docx.Document()
    d.sections[0].left_margin = Pt(72)
    d.add_heading("Worksheet", level=1)
    d.add_paragraph("(a) Explain what an IP address is.")
    d.add_paragraph("")  # blank answer paragraph -> index 2
    d.add_paragraph("(b) State the number of bits in IPv6.")
    d.add_paragraph("")  # blank answer paragraph -> index 4
    table = d.add_table(rows=3, cols=3)
    table.rows[0].cells[0].text = "Statement"
    table.rows[0].cells[1].text = "True"
    table.rows[0].cells[2].text = "False"
    table.rows[1].cells[0].text = "IPv6 is 128-bit."
    table.rows[2].cells[0].text = "MAC means machine access control."
    p = tmp_path / "ws.docx"
    d.save(str(p))
    return p


def test_extract_structure(worksheet_docx: Path):
    s = docx_backend.extract_structure(str(worksheet_docx))
    # blank paragraphs are included (they are answer slots)
    blanks = [p for p in s["paragraphs"] if p["text"] == ""]
    assert len(blanks) >= 2
    assert len(s["tables"]) == 1
    assert s["tables"][0]["rows"][0] == ["Statement", "True", "False"]


def test_apply_ops_in_place_and_preserved(worksheet_docx: Path, tmp_path: Path):
    import shutil

    original = tmp_path / "orig.docx"
    shutil.copyfile(str(worksheet_docx), str(original))

    ops = [
        {"op": "set_paragraph", "index": 2, "text": "A unique address for a device on a network."},
        {"op": "set_paragraph", "index": 4, "text": "128"},
        {"op": "set_cell", "table": 0, "row": 1, "col": 1, "text": "✓"},  # True
        {"op": "set_cell", "table": 0, "row": 2, "col": 2, "text": "✓"},  # False
    ]
    result = docx_backend.apply_ops(str(worksheet_docx), ops)
    assert result.in_place is True
    assert result.answers_written == 4

    d = docx.Document(str(worksheet_docx))
    assert d.paragraphs[2].text == "A unique address for a device on a network."
    assert d.paragraphs[4].text == "128"
    # question paragraphs untouched
    assert d.paragraphs[1].text == "(a) Explain what an IP address is."
    tb = d.tables[0]
    assert tb.rows[1].cells[1].text == "✓"  # IPv6 128-bit -> True
    assert tb.rows[2].cells[2].text == "✓"  # MAC wrong name -> False
    # section margins preserved
    assert d.sections[0].left_margin == pytest.approx(914400, abs=1)


def test_styles_preserved(worksheet_docx: Path, tmp_path: Path):
    import shutil
    from lxml import etree

    original = tmp_path / "orig.docx"
    shutil.copyfile(str(worksheet_docx), str(original))
    ops = [{"op": "set_paragraph", "index": 2, "text": "x"}]
    docx_backend.apply_ops(str(worksheet_docx), ops)

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    def style_count(p):
        with zipfile.ZipFile(p) as z:
            x = etree.fromstring(z.read("word/styles.xml"))
        return len(x.findall(".//w:style", ns))

    assert style_count(original) == style_count(worksheet_docx)


def test_extract_json_array_variants():
    assert extract_json_array('[{"op":"set_paragraph","index":1,"text":"a"}]')[0]["index"] == 1
    fenced = "```json\n[{\"op\":\"set_cell\",\"table\":0,\"row\":1,\"col\":1,\"text\":\"✓\"}]\n```"
    assert extract_json_array(fenced)[0]["op"] == "set_cell"
    wrapped = '{"operations": [{"op":"set_paragraph","index":2,"text":"b"}]}'
    assert extract_json_array(wrapped)[0]["index"] == 2
