"""Phase 5 browser-level acceptance: drive the real responsive web UI.

These tests launch the actual server and a headless Chromium, exercise the
UI on both a desktop and a mobile viewport, and assert the UI shows the
correct results and the equivalent CLI command. This proves the GUI works
and stays a thin, equivalent layer over the CLI -- no manual QA needed.
"""
from __future__ import annotations

import socket
import subprocess
import sys
import time
from pathlib import Path

import httpx
import pytest

playwright_sync = pytest.importorskip("playwright.sync_api")
from playwright.sync_api import sync_playwright  # noqa: E402


def _free_port() -> int:
    s = socket.socket()
    s.bind(("127.0.0.1", 0))
    port = s.getsockname()[1]
    s.close()
    return port


@pytest.fixture(scope="module")
def server():
    port = _free_port()
    proc = subprocess.Popen(
        [sys.executable, "-m", "uvicorn", "autodoc.api:app",
         "--host", "127.0.0.1", "--port", str(port), "--log-level", "warning"],
    )
    base = f"http://127.0.0.1:{port}"
    deadline = time.time() + 30
    ok = False
    while time.time() < deadline:
        try:
            if httpx.get(base + "/api/version", timeout=1).status_code == 200:
                ok = True
                break
        except Exception:
            time.sleep(0.3)
    if not ok:
        proc.terminate()
        pytest.fail("server did not start in time")
    yield base
    proc.terminate()
    try:
        proc.wait(timeout=10)
    except Exception:
        proc.kill()


@pytest.fixture
def quiz_file(tmp_path: Path) -> Path:
    import docx

    d = docx.Document()
    d.add_heading("测验", level=1)
    d.add_paragraph("1. 中国的首都是____。")
    d.add_paragraph("2. 请简述水的化学式。")
    p = tmp_path / "quiz.docx"
    d.save(str(p))
    return p


VIEWPORTS = {
    "desktop": {"width": 1280, "height": 900},
    "mobile": {"width": 390, "height": 844},  # iPhone 12-ish
}


@pytest.mark.parametrize("viewport_name", ["desktop", "mobile"])
def test_gui_parse_and_edit(server, quiz_file, viewport_name, tmp_path):
    with sync_playwright() as pw:
        browser = pw.chromium.launch(headless=True)
        context = browser.new_context(viewport=VIEWPORTS[viewport_name])
        page = context.new_page()
        page.goto(server, wait_until="networkidle")

        # meta loaded
        page.wait_for_function("document.getElementById('version').textContent !== '…'")
        assert ".docx" in page.get_by_test_id("formats").inner_text()

        # upload + parse
        page.get_by_test_id("file").set_input_files(str(quiz_file))
        page.get_by_test_id("btn-parse").click()
        page.wait_for_selector("#questions li")
        questions = page.locator("#questions li")
        assert questions.count() == 2
        assert "autodoc parse" in page.get_by_test_id("cli-parse").inner_text()

        # edit -> triggers download, shows equivalent CLI
        page.get_by_test_id("answers").fill('{"q1":"北京","q2":"H2O"}')
        with page.expect_download() as dl_info:
            page.get_by_test_id("btn-edit").click()
        download = dl_info.value
        out = tmp_path / f"{viewport_name}_out.docx"
        download.save_as(str(out))

        cli_text = page.get_by_test_id("cli-edit").inner_text()
        assert "autodoc edit" in cli_text
        assert "answers_written=2" in cli_text

        # the downloaded file really contains the answers, in-place style
        import docx

        full = "\n".join(p.text for p in docx.Document(str(out)).paragraphs)
        assert "北京" in full and "H2O" in full
        assert "测验" in full  # heading preserved

        context.close()
        browser.close()
