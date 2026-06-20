"""Opt-in live end-to-end test against a real OpenAI-compatible model.

Skipped by default. To run it against a real endpoint::

    set AUTODOC_E2E=1
    set AUTODOC_BASE_URL=https://api.deepseek.com
    set AUTODOC_MODEL=deepseek-v4-flash
    set AUTODOC_API_KEY=sk-...
    pytest -q tests/test_e2e_live.py

Requires the account to have balance; otherwise the provider returns an
error (e.g. 402 Payment Required) and the test fails with that message.
"""
from __future__ import annotations

import os
from pathlib import Path

import docx
import pytest

pytestmark = pytest.mark.skipif(
    os.environ.get("AUTODOC_E2E") != "1",
    reason="live e2e disabled; set AUTODOC_E2E=1 and AUTODOC_* env vars to run",
)


def test_live_answer_docx(tmp_path: Path):
    from autodoc.core.agent import answer_document
    from autodoc.core.llm import OpenAICompatClient

    d = docx.Document()
    d.add_paragraph("1. 中国的首都是____。")
    d.add_paragraph("2. 水的化学式是____。")
    path = tmp_path / "live.docx"
    d.save(str(path))

    client = OpenAICompatClient(
        base_url=os.environ.get("AUTODOC_BASE_URL", ""),
        model=os.environ.get("AUTODOC_MODEL", ""),
        api_key=os.environ.get("AUTODOC_API_KEY", ""),
    )
    result = answer_document(str(path), client)
    assert result.answers_written >= 1
    full = "\n".join(p.text for p in docx.Document(str(path)).paragraphs)
    assert "____" not in full  # blanks got filled
