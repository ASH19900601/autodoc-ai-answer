"""Phase 1+2 tests: DOCX parsing, in-place editing, format preservation."""
from __future__ import annotations

from pathlib import Path

import docx
import pytest

from autodoc.core.edit import write_answers
from autodoc.core.models import Answer, DocFormat, QuestionType
from autodoc.core.parse import parse_document
from conftest import read_zip_member, zip_members


# ---------- Phase 1: parsing ----------

def test_parse_finds_questions(sample_docx: Path):
    result = parse_document(str(sample_docx))
    assert result.format == DocFormat.docx
    numbers = [q.number for q in result.questions]
    assert numbers == ["1", "2", "3"]


def test_parse_question_types(sample_docx: Path):
    result = parse_document(str(sample_docx))
    by_num = {q.number: q for q in result.questions}
    assert by_num["1"].qtype == QuestionType.fill_blank
    assert by_num["2"].qtype == QuestionType.short_answer
    assert by_num["3"].qtype == QuestionType.fill_blank


def test_non_question_paragraph_ignored(sample_docx: Path):
    result = parse_document(str(sample_docx))
    texts = [q.text for q in result.questions]
    assert not any("说明" in t for t in texts)


# ---------- Phase 2: in-place editing ----------

def test_edit_writes_answers_in_place(sample_docx: Path):
    questions = parse_document(str(sample_docx)).questions
    answers = [
        Answer(question_id="q1", text="北京"),
        Answer(question_id="q2", text="H2O"),
        Answer(question_id="q3", text="月球"),
    ]
    result = write_answers(str(sample_docx), answers, questions)
    assert result.in_place is True
    assert result.path == str(sample_docx)
    assert result.answers_written == 3


def test_answers_present_after_edit(sample_docx: Path):
    questions = parse_document(str(sample_docx)).questions
    answers = [
        Answer(question_id="q1", text="北京"),
        Answer(question_id="q2", text="H2O"),
        Answer(question_id="q3", text="月球"),
    ]
    write_answers(str(sample_docx), answers, questions)

    document = docx.Document(str(sample_docx))
    full = "\n".join(p.text for p in document.paragraphs)
    assert "北京" in full
    assert "H2O" in full
    assert "月球" in full
    # blank markers must be gone where we replaced them
    assert "____" not in "1. 中国的首都是北京。"  # sanity
    q1_para = [p for p in document.paragraphs if p.text.startswith("1.")][0]
    assert "____" not in q1_para.text
    assert q1_para.text == "1. 中国的首都是北京。"


def test_question_count_unchanged_after_edit(sample_docx: Path):
    before = len(parse_document(str(sample_docx)).questions)
    questions = parse_document(str(sample_docx)).questions
    answers = [Answer(question_id=q.id, text="X") for q in questions]
    write_answers(str(sample_docx), answers, questions)
    after = len(parse_document(str(sample_docx)).questions)
    assert before == after == 3


# ---------- Format preservation (strict, automatable) ----------

PRESERVED_MEMBERS = [
    "word/styles.xml",
    "word/theme/theme1.xml",
    "word/settings.xml",
    "word/fontTable.xml",
    "[Content_Types].xml",
    "word/numbering.xml",
]


def test_format_preserved_byte_identical(sample_docx: Path, tmp_path: Path):
    original = tmp_path / "original.docx"
    import shutil

    shutil.copyfile(str(sample_docx), str(original))

    questions = parse_document(str(sample_docx)).questions
    answers = [Answer(question_id=q.id, text="答案") for q in questions]
    write_answers(str(sample_docx), answers, questions)

    # the zip member set is unchanged (nothing added/removed)
    assert zip_members(original) == zip_members(sample_docx)

    # style / theme / settings / numbering members are byte-identical
    for member in PRESERVED_MEMBERS:
        assert read_zip_member(original, member) == read_zip_member(
            sample_docx, member
        ), f"{member} changed during edit"

    # document.xml MUST differ (we wrote answers into the body)
    assert read_zip_member(original, "word/document.xml") != read_zip_member(
        sample_docx, "word/document.xml"
    )


def test_section_margins_preserved(sample_docx: Path):
    questions = parse_document(str(sample_docx)).questions
    answers = [Answer(question_id=q.id, text="答案") for q in questions]
    write_answers(str(sample_docx), answers, questions)
    document = docx.Document(str(sample_docx))
    section = document.sections[0]
    assert section.left_margin == pytest.approx(914400, abs=1)  # 72pt in EMU
    assert section.right_margin == pytest.approx(914400, abs=1)


def test_bold_blank_replacement_keeps_bold(sample_docx: Path):
    questions = parse_document(str(sample_docx)).questions
    answers = [Answer(question_id="q3", text="月球")]
    write_answers(str(sample_docx), answers, questions)
    document = docx.Document(str(sample_docx))
    q3 = [p for p in document.paragraphs if p.text.startswith("3.")][0]
    # the run that used to hold the bold blank now holds the answer, still bold
    bold_runs = [r for r in q3.runs if r.bold and "月球" in r.text]
    assert bold_runs, "answer should inherit the bold formatting of the blank run"
