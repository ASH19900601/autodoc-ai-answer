"""Deterministic tests for the expanded worksheet op set (all question types)."""
from __future__ import annotations

from pathlib import Path

import docx
import pytest

from autodoc.core.answer_router import choose_mode
from autodoc.core.backends import docx_backend


def _texts(path):
    return [p.text for p in docx.Document(str(path)).paragraphs]


def test_fill_blank_in_paragraph_keeps_other_text(tmp_path: Path):
    d = docx.Document()
    d.add_paragraph("1. 中国的首都是____，最大的城市是____。")
    p = tmp_path / "a.docx"
    d.save(str(p))

    ops = [
        {"op": "fill_blank", "index": 0, "occurrence": 1, "text": "北京"},
        {"op": "fill_blank", "index": 0, "occurrence": 2, "text": "上海"},
    ]
    r = docx_backend.apply_ops(str(p), ops)
    assert r.answers_written == 2
    assert _texts(p)[0] == "1. 中国的首都是北京，最大的城市是上海。"


def test_fill_blank_missing_occurrence_is_skipped(tmp_path: Path):
    d = docx.Document()
    d.add_paragraph("只有一个空：____")
    p = tmp_path / "b.docx"
    d.save(str(p))
    ops = [{"op": "fill_blank", "index": 0, "occurrence": 5, "text": "x"}]
    r = docx_backend.apply_ops(str(p), ops)
    assert r.answers_written == 0
    assert _texts(p)[0] == "只有一个空：____"


def test_append_paragraph_annotates_mc_option(tmp_path: Path):
    d = docx.Document()
    d.add_paragraph("2. 下列哪个是质数？")
    d.add_paragraph("A. 4")
    d.add_paragraph("B. 7")
    p = tmp_path / "mc.docx"
    d.save(str(p))
    ops = [{"op": "append_paragraph", "index": 2, "text": "✓(答案:B)"}]
    docx_backend.apply_ops(str(p), ops)
    assert "✓(答案:B)" in _texts(p)[2]
    assert _texts(p)[1] == "A. 4"  # other option untouched


def test_append_cell(tmp_path: Path):
    d = docx.Document()
    t = d.add_table(rows=2, cols=2)
    t.rows[0].cells[0].text = "国家"
    t.rows[0].cells[1].text = "首都"
    t.rows[1].cells[0].text = "中国"
    t.rows[1].cells[1].text = ""
    p = tmp_path / "tbl.docx"
    d.save(str(p))
    ops = [{"op": "set_cell", "table": 0, "row": 1, "col": 1, "text": "北京"}]
    docx_backend.apply_ops(str(p), ops)
    d2 = docx.Document(str(p))
    assert d2.tables[0].rows[1].cells[1].text == "北京"
    assert d2.tables[0].rows[0].cells[0].text == "国家"  # header untouched


def test_choice_doc_routes_to_worksheet(tmp_path: Path):
    d = docx.Document()
    d.add_paragraph("1. 下列哪个是质数？")
    d.add_paragraph("A. 4")
    d.add_paragraph("B. 7")
    d.add_paragraph("C. 9")
    p = tmp_path / "choice.docx"
    d.save(str(p))
    assert choose_mode(str(p)) == "worksheet"


def test_cloze_multiblank_routes_to_worksheet(tmp_path: Path):
    d = docx.Document()
    d.add_paragraph("1. 中国的首都是____，水的化学式是____。")
    p = tmp_path / "cloze.docx"
    d.save(str(p))
    assert choose_mode(str(p)) == "worksheet"


def test_single_blank_stays_simple(tmp_path: Path):
    d = docx.Document()
    d.add_paragraph("1. 中国的首都是____。")
    p = tmp_path / "single.docx"
    d.save(str(p))
    assert choose_mode(str(p)) == "simple"
