"""Generate a comprehensive sample quiz covering all question types.

Produces ``examples/sample-all-types.docx`` with: fill-in-the-blank (incl.
cloze/multi-blank), single-choice, multi-choice, true/false (inline and
tick-table), short answer, essay, matching (table), ordering, and
calculation. Useful for manually exercising worksheet mode end-to-end.

    python examples/make_sample_quiz.py
"""
from __future__ import annotations

import os

import docx
from docx.shared import Pt


def build(path: str) -> str:
    d = docx.Document()
    d.sections[0].left_margin = Pt(72)
    d.add_heading("综合题型测试卷", level=1)
    d.add_paragraph("说明：请在横线、括号、空白处或表格中作答。")

    d.add_heading("一、填空题", level=2)
    d.add_paragraph("1. 中国的首都是____，水的化学式是____。")

    d.add_heading("二、单选题", level=2)
    d.add_paragraph("2. 下列哪个数是质数？")
    d.add_paragraph("A. 4")
    d.add_paragraph("B. 6")
    d.add_paragraph("C. 7")
    d.add_paragraph("D. 8")
    d.add_paragraph("答案：____")

    d.add_heading("三、多选题", level=2)
    d.add_paragraph("3. 下列属于编程语言的有（多选）？")
    d.add_paragraph("A. Python")
    d.add_paragraph("B. HTML")
    d.add_paragraph("C. Java")
    d.add_paragraph("D. Microsoft Word")
    d.add_paragraph("答案：____")

    d.add_heading("四、判断题（行内）", level=2)
    d.add_paragraph("4. 地球是太阳系中体积最大的行星。（____）")

    d.add_heading("五、判断题（表格打勾）", level=2)
    t = d.add_table(rows=3, cols=3)
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "陈述", "正确", "错误"
    t.rows[1].cells[0].text = "HTTP 默认使用 80 端口。"
    t.rows[2].cells[0].text = "一个字节等于 4 个比特。"

    d.add_heading("六、简答题", level=2)
    d.add_paragraph("6. 简述什么是闭包（closure）。")
    d.add_paragraph("")  # answer slot

    d.add_heading("七、论述题", level=2)
    d.add_paragraph("7. 论述 HTTP 与 HTTPS 的主要区别。")
    d.add_paragraph("")  # answer slot

    d.add_heading("八、匹配题", level=2)
    d.add_paragraph("8. 将下列国家与其首都匹配，在右列填入对应字母：A. 东京  B. 北京  C. 巴黎")
    m = d.add_table(rows=4, cols=2)
    m.style = "Table Grid"
    m.rows[0].cells[0].text = "国家"
    m.rows[0].cells[1].text = "首都（填字母）"
    m.rows[1].cells[0].text = "中国"
    m.rows[2].cells[0].text = "日本"
    m.rows[3].cells[0].text = "法国"

    d.add_heading("九、排序题", level=2)
    d.add_paragraph("9. 将以下步骤按正确顺序排列（如 2-1-3-4）：①编译 ②编写代码 ③运行 ④需求分析")
    d.add_paragraph("答案：____")

    d.add_heading("十、计算题", level=2)
    d.add_paragraph("10. 计算 12 × 8 = ____。")

    d.add_paragraph("[Total 10 questions]")
    d.save(path)
    return path


if __name__ == "__main__":
    here = os.path.dirname(os.path.abspath(__file__))
    out = build(os.path.join(here, "sample-all-types.docx"))
    print("wrote", out)
