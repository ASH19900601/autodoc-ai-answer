"""DOCX backend: parse questions and write answers back in-place.

Format preservation strategy
----------------------------
We load the existing .docx with python-docx and save it back to the SAME
path. python-docx round-trips everything it does not touch (styles,
numbering, themes, section properties, images, ...). We only either:

* replace a blank (runs of ``_`` / ``＿``) inside the question paragraph, or
* append an answer run to the end of the question paragraph,

copying the formatting of an existing run so the inserted text matches the
surrounding style. We never rebuild the document, so layout is preserved.
"""
from __future__ import annotations

import json
import re
from typing import Dict, List, Tuple

import docx  # python-docx

from ..models import Anchor, Answer, DocFormat, EditResult, Question, QuestionType

# A question paragraph: starts with a number then . 、 ) ） followed by text.
_QUESTION_RE = re.compile(r"^\s*(\d+)\s*[.、)）]\s*(\S.*)$")
# A blank to fill: two or more ASCII or full-width underscores.
_BLANK_RE = re.compile(r"[_＿]{2,}")


def _para_text(paragraph) -> str:
    return "".join(run.text for run in paragraph.runs)


def parse_questions(path: str) -> List[Question]:
    document = docx.Document(path)
    questions: List[Question] = []
    for idx, paragraph in enumerate(document.paragraphs):
        text = _para_text(paragraph)
        m = _QUESTION_RE.match(text)
        if not m:
            continue
        number = m.group(1)
        blank = _BLANK_RE.search(text)
        if blank:
            mode = "replace_blank"
            qtype = QuestionType.fill_blank
            detail = {"mode": mode}
        else:
            mode = "append"
            qtype = QuestionType.short_answer
            detail = {"mode": mode}
        locator = json.dumps({"para_index": idx, "mode": mode})
        questions.append(
            Question(
                id=f"q{number}",
                number=number,
                text=text.strip(),
                qtype=qtype,
                anchor=Anchor(
                    format=DocFormat.docx, locator=locator, detail=detail
                ),
            )
        )
    return questions


def _run_map(paragraph) -> List[Tuple[int, int, object]]:
    """Return [(start, end, run), ...] over the concatenated paragraph text."""
    spans = []
    pos = 0
    for run in paragraph.runs:
        length = len(run.text)
        spans.append((pos, pos + length, run))
        pos += length
    return spans


def _replace_range(paragraph, start: int, end: int, new_text: str) -> None:
    """Replace text[start:end] in a paragraph, preserving run formatting.

    The replacement text is placed into the run that contains ``start``;
    characters of the deleted range that fall into later runs are removed.
    """
    spans = _run_map(paragraph)
    for s, e, run in spans:
        if e <= start or s >= end:
            continue  # run not touched
        local_start = max(start, s) - s
        local_end = min(end, e) - s
        if s <= start < e:
            # first affected run: splice in the replacement
            run.text = run.text[:local_start] + new_text + run.text[local_end:]
        else:
            # subsequent affected runs: drop the overlapping chars
            run.text = run.text[:local_start] + run.text[local_end:]


def _append_answer(paragraph, answer: str) -> None:
    sep = "" if _para_text(paragraph).endswith((" ", "：", ":")) else " "
    if paragraph.runs:
        src = paragraph.runs[-1]
        new_run = paragraph.add_run(sep + answer)
        # copy basic character formatting from the last run
        new_run.bold = src.bold
        new_run.italic = src.italic
        new_run.underline = src.underline
        if src.font is not None and src.font.name:
            new_run.font.name = src.font.name
        if src.font is not None and src.font.size:
            new_run.font.size = src.font.size
    else:
        paragraph.add_run(sep + answer)


def write_answers(
    path: str, answers: List[Answer], questions: List[Question], output: str = None
) -> EditResult:
    document = docx.Document(path)
    paragraphs = document.paragraphs
    by_id: Dict[str, Answer] = {a.question_id: a for a in answers}
    written = 0

    for q in questions:
        ans = by_id.get(q.id)
        if ans is None:
            continue
        loc = json.loads(q.anchor.locator)
        para = paragraphs[loc["para_index"]]
        if loc["mode"] == "replace_blank":
            text = _para_text(para)
            blank = _BLANK_RE.search(text)
            if blank:
                _replace_range(para, blank.start(), blank.end(), ans.text)
            else:
                _append_answer(para, ans.text)
        else:
            _append_answer(para, ans.text)
        written += 1

    target = output or path
    document.save(target)
    return EditResult(
        path=target,
        format=DocFormat.docx,
        answers_written=written,
        in_place=(target == path),
    )


# ===================================================================
# Worksheet mode: structure extraction + in-place edit operations.
# Used for real-world worksheets that don't follow the "1. ____" format
# (lettered sub-parts, written-answer paragraphs, tick-box tables).
# ===================================================================

def extract_structure(path: str) -> dict:
    """Return paragraphs (index/style/text, incl. blanks) and tables (cells)."""
    document = docx.Document(path)
    paragraphs = [
        {"index": i, "style": p.style.name if p.style else "", "text": _para_text(p)}
        for i, p in enumerate(document.paragraphs)
    ]
    tables = []
    for ti, tb in enumerate(document.tables):
        rows = []
        for ri, row in enumerate(tb.rows):
            rows.append([cell.text for cell in row.cells])
        tables.append({"index": ti, "rows": rows})
    return {"paragraphs": paragraphs, "tables": tables}


def _set_paragraph_text(paragraph, text: str) -> None:
    """Set a paragraph's text in place, preserving its style.

    If it already has runs, reuse the first run (keeps char formatting) and
    clear the rest. If it is empty, add a run (paragraph style is preserved).
    """
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def apply_ops(path: str, ops: list, output: str = None) -> EditResult:
    """Apply worksheet edit operations in place.

    Supported ops:
      {"op": "set_paragraph", "index": int, "text": str}
      {"op": "append_paragraph", "index": int, "text": str}
      {"op": "set_cell", "table": int, "row": int, "col": int, "text": str}
    """
    document = docx.Document(path)
    paragraphs = document.paragraphs
    tables = document.tables
    written = 0
    for op in ops:
        kind = op.get("op")
        try:
            if kind == "set_paragraph":
                _set_paragraph_text(paragraphs[op["index"]], str(op["text"]))
                written += 1
            elif kind == "append_paragraph":
                _append_answer(paragraphs[op["index"]], str(op["text"]))
                written += 1
            elif kind == "set_cell":
                cell = tables[op["table"]].rows[op["row"]].cells[op["col"]]
                para = cell.paragraphs[0]
                _set_paragraph_text(para, str(op["text"]))
                written += 1
        except (IndexError, KeyError):
            continue
    target = output or path
    document.save(target)
    return EditResult(
        path=target,
        format=DocFormat.docx,
        answers_written=written,
        in_place=(target == path),
    )
